from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from .forms import BookForm
from .models import Book
from docx import Document
import re

def homePage(request):
    return render(request, 'home.html')

# Create your views here.
def upload(request):
    context = {}
    if request.method == "POST":
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        name = fs.save(uploaded_file.name, uploaded_file)
        context['url'] = fs.url(name)
        return render(request, 'upload.html', context)
    return render(request, 'upload.html')

def book_list(request):
    books= Book.objects.all()
    return render(request, 'book_list.html', {'books':books})

def upload_book(request):
    if request.method == 'POST':
        document = Document(request.FILES['document'])
        for i in range(len(document.paragraphs)):
           document.paragraphs[i].text = re.sub(r"account", "SANITIZED", document.paragraphs[i].text)
        document.save(request.FILES['document'])
        form = BookForm(request.POST, request.FILES)
        if form.is_valid():
            name = form.save()
            print (name)
            return redirect('book_list')
    else:
        form = BookForm()
    return render(request, 'upload_book.html', {'form': form})
